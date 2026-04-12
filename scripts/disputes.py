"""Credify dispute engine — 3-gate classification, letter generation, FTC summary.
Codified from the dispute-letters SKILL.md."""
import os
import stat
from datetime import datetime
from classify import parse_balance, parse_open_closed, resolve_client_name, safe_filename

# Bureau addresses
BUREAU_ADDRESSES = {
    "EXP": {"name": "Experian", "address": "P.O. Box 4500\nAllen, TX 75013"},
    "TU": {"name": "TransUnion", "address": "P.O. Box 2000\nChester, PA 19016"},
    "EQF": {"name": "Equifax", "address": "P.O. Box 740256\nAtlanta, GA 30374-0256"},
}

LEGAL_BODY = """These accounts are fraudulent and the direct result of identity theft tied to the Equifax data breach of 2017, which compromised the sensitive personal data of more than 147 million Americans, including Social Security numbers, credit card details, and other critical identifiers.

It is well-documented that this breach has led to widespread fraudulent activity, with criminals using stolen data to open unauthorized credit accounts. The accounts appearing on my report fit this exact pattern. I have never authorized these accounts, and any claim otherwise must be substantiated with legally admissible documentation --- including signed applications, IP addresses, and identity verification records.

Under the Fair Credit Reporting Act (FCRA), Section 611 (15 U.S.C. \u00a71681i), I am exercising my full right to dispute inaccurate, incomplete, or unverifiable information. Federal law requires you to conduct a reasonable investigation within 30 days and either verify the accounts with proper documentation or delete them entirely. Simply "parroting" information from the furnisher without independent verification is a violation of the FCRA, as confirmed in numerous CFPB enforcement actions.

Further, I draw your attention to the Consumer Financial Protection Bureau's lawsuit against Experian, which alleges systemic violations including the failure to properly investigate disputes and the continued reporting of unverifiable data. These violations directly mirror what occurs when fraudulent accounts from a known breach remain on a consumer's file. Be advised: any failure on your part to perform a thorough investigation will be documented and reported to the CFPB, the Federal Trade Commission (FTC), and my state Attorney General for enforcement.

Let me be perfectly clear:

- If you cannot produce hard proof --- not assumptions, not hearsay, but verifiable evidence that I personally opened and authorized these accounts --- you are legally obligated to delete them.

- Continuing to report fraudulent accounts in light of the Equifax breach constitutes willful noncompliance with federal law and subjects your agency to liability under FCRA \u00a7616 and \u00a7617, which provide for actual, statutory, and punitive damages, plus attorney's fees.

This is my FINAL NOTICE. I have previously disputed these items and they remain on my credit file without proper verification. This letter places you on notice that should you fail to remove the fraudulent accounts or attempt to verify them without legitimate documentation, I will pursue all available remedies, including litigation and formal complaints to the CFPB, FTC, and my state Attorney General. I expect written confirmation of deletion within the time allowed by law. Failure to comply will result in immediate escalation without further warning."""

MAX_LETTERS = 50


def classify_disputes(derog):
    """3-gate decision tree. Returns (standard_list, isolated_list)."""
    standard, isolated = [], []
    for a in derog:
        bureaus = a.get("bureaus", [])
        # Gate 1: single-bureau — no asymmetry possible
        if len(bureaus) <= 1:
            standard.append(a)
            continue
        # Gate 2: toxic types always standard batch
        if a.get("is_collection") or a.get("is_chargeoff") or a.get("is_repo"):
            standard.append(a)
            continue
        # Gate 2b: derogatory on ALL bureaus — no asymmetry
        # (for now, treat all multi-bureau derog as standard — isolation requires
        # per-bureau derogatory status which parsers don't yet track individually)
        standard.append(a)
    return standard, isolated


def build_shared_blocks(client, inquiries=None, addresses=None, name_variations=None):
    """Pre-assemble shared content blocks. Handles missing data gracefully."""
    inquiries = inquiries or []
    addresses = addresses or []
    name_variations = name_variations or []
    client_name = resolve_client_name(client).upper()
    return {
        "client_name": client_name or "CLIENT",
        "client_address": client.get("address", ""),
        "inquiries": inquiries,
        "addresses": addresses,
        "name_variations": name_variations,
    }


def generate_disputes(data, derog, output_dir):
    """Public orchestrator. Returns list of (docx_path, recipient_dict) tuples."""
    if not derog:
        return []
    os.makedirs(output_dir, mode=0o700, exist_ok=True)
    client = data.get("client", {})
    shared = build_shared_blocks(
        client,
        data.get("inquiries", []),
        data.get("addresses", []),
        data.get("name_variations", []),
    )
    standard, isolated = classify_disputes(derog)
    results = []
    letter_count = 0
    # Batch standard: 5 per letter
    for i in range(0, len(standard), 5):
        if letter_count >= MAX_LETTERS:
            break
        batch = standard[i:i+5]
        letter_count += 1
        filename = f"Dispute_{letter_count}.docx"
        path = _generate_letter(client, batch, shared, list(BUREAU_ADDRESSES.keys()), filename, output_dir)
        # Standard letters go to all 3 bureaus
        for bureau_code, bureau_info in BUREAU_ADDRESSES.items():
            results.append((path, bureau_info))
    # Isolated: 1 per letter (future — when parsers track per-bureau derog status)
    for a in isolated:
        if letter_count >= MAX_LETTERS:
            break
        letter_count += 1
        target_bureaus = a.get("derog_bureaus", a.get("bureaus", []))
        creditor_safe = a.get("creditor", "Unknown").replace("/", "-").replace(" ", "_")[:20]
        filename = f"Isolated_{creditor_safe}_{'+'.join(target_bureaus)}.docx"
        path = _generate_letter(client, [a], shared, target_bureaus, filename, output_dir)
        for bc in target_bureaus:
            if bc in BUREAU_ADDRESSES:
                results.append((path, BUREAU_ADDRESSES[bc]))
    # FTC Summary
    ftc_path = _generate_ftc_summary(client, derog, shared, output_dir)
    if ftc_path:
        results.append((ftc_path, {"name": "FTC Summary", "address": ""}))
    return results


def _generate_letter(client, accounts, shared, bureau_codes, filename, output_dir):
    """Generate a single dispute letter .docx."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Client header
    doc.add_paragraph(shared["client_name"])
    if shared["client_address"]:
        doc.add_paragraph(shared["client_address"])
    doc.add_paragraph("")

    # Bureau addresses
    for bc in bureau_codes:
        if bc in BUREAU_ADDRESSES:
            info = BUREAU_ADDRESSES[bc]
            doc.add_paragraph(f"{info['name']}\n{info['address']}")
    doc.add_paragraph("")

    # Subject
    p = doc.add_paragraph()
    run = p.add_run("Subject: FINAL NOTICE -- Equifax Data Breach Dispute -- Demand for Immediate Deletion of Fraudulent Accounts Based on CFPB Action")
    run.bold = True
    doc.add_paragraph("")
    doc.add_paragraph("To Whom It May Concern,")
    doc.add_paragraph("")
    doc.add_paragraph("I am formally disputing the presence of the following accounts on my credit file, which I believe are the result of identity theft stemming from the Equifax data breach:")
    doc.add_paragraph("")

    # Accounts
    for i, a in enumerate(accounts, 1):
        balance = parse_balance(a.get("balance", "0"))
        p = doc.add_paragraph()
        acct_num = a.get('account_number', '')
        acct_str = f", {acct_num}" if acct_num else ""
        run = p.add_run(f"{i}. {a.get('creditor', 'Unknown')}{acct_str}, BALANCE: ${balance:,.0f}")
        run.bold = True
    doc.add_paragraph("")

    # Addresses section
    addrs = shared.get("addresses", [])
    if addrs:
        p = doc.add_paragraph()
        run = p.add_run("FRAUDULENT ADDRESSES -- DELETE THESE IMMEDIATELY:")
        run.bold = True
        run.underline = True
        for i, addr in enumerate(addrs, 1):
            addr_text = addr if isinstance(addr, str) else addr.get("address", "")
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {addr_text}")
            run.bold = True
        doc.add_paragraph("")

    # Inquiries section
    inqs = shared.get("inquiries", [])
    if inqs:
        p = doc.add_paragraph()
        run = p.add_run("FRAUDULENT INQUIRIES -- DELETE THESE IMMEDIATELY:")
        run.bold = True
        run.underline = True
        for i, inq in enumerate(inqs, 1):
            if isinstance(inq, str):
                text = inq
            else:
                text = f"{inq.get('creditor', '?')} Inquired on {inq.get('date', '?')}"
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {text}")
            run.bold = True
        doc.add_paragraph("")

    # Name variations
    names = shared.get("name_variations", [])
    if names:
        p = doc.add_paragraph()
        run = p.add_run(f"FRAUDULENT ITERATION OF MY NAME: My only name is {shared['client_name']}. Delete these:")
        run.bold = True
        run.underline = True
        for i, name in enumerate(names, 1):
            name_text = name if isinstance(name, str) else name.get("name", "")
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {name_text}")
            run.bold = True
        doc.add_paragraph("")

    # Legal body
    doc.add_paragraph(LEGAL_BODY)
    doc.add_paragraph("")
    doc.add_paragraph("Thank you for your prompt attention to this matter.")
    doc.add_paragraph("")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(shared["client_name"])

    # Save with restrictive permissions
    path = os.path.join(output_dir, filename)
    doc.save(path)
    os.chmod(path, stat.S_IRUSR | stat.S_IWUSR)
    return path


def _generate_ftc_summary(client, derog, shared, output_dir):
    """Generate FTC Identity Theft Report summary .docx."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    p = doc.add_paragraph()
    run = p.add_run("FTC Identity Theft Report \u2014 Account Summary")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph("")

    # Client info
    p = doc.add_paragraph()
    run = p.add_run(f"Client: {shared['client_name']}")
    run.bold = True
    if client.get("dob_year"):
        p = doc.add_paragraph()
        run = p.add_run(f"Year of Birth: {client['dob_year']}")
        run.bold = True
    if shared["client_address"]:
        p = doc.add_paragraph()
        run = p.add_run(f"Address: {shared['client_address']}")
        run.bold = True
    p = doc.add_paragraph()
    run = p.add_run(f"Date: {datetime.now().strftime('%m/%d/%Y')}")
    run.bold = True
    doc.add_paragraph("")

    # Table
    sorted_derog = sorted(derog, key=lambda a: a.get("opened", ""))
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    headers = ["#", "Creditor Name", "Account Number", "Date Opened", "Balance", "Status", "Bureau(s)"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for i, a in enumerate(sorted_derog, 1):
        row = table.add_row().cells
        balance = parse_balance(a.get("balance", "0"))
        bureaus = a.get("bureaus", [])
        status = a.get("payment_status", a.get("status", ""))
        row[0].text = str(i)
        row[1].text = a.get("creditor", "Unknown")
        row[2].text = a.get("account_number", "")
        row[3].text = a.get("opened", "")
        row[4].text = f"${balance:,.0f}"
        row[5].text = status[:30] if status else ""
        row[6].text = ", ".join(bureaus)

    filename = f"{safe_filename(resolve_client_name(client))}_FTC_Summary.docx"
    path = os.path.join(output_dir, filename)
    doc.save(path)
    os.chmod(path, stat.S_IRUSR | stat.S_IWUSR)
    return path
